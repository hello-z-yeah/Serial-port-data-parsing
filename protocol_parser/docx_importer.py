"""从 Word 协议文档自动生成 JSON 配置。

支持的文档结构：
- 表格型：识别包含 "命令字"/"cmd"/"attrid" 等关键词的表格
- 段落型：提取关键信息（帧头、版本、校验等）
- 混合型：自动适配

重要策略：
- 以 V3.0 串口标准协议（v3_serial.json）作为基底
- 任何导入的新协议仅在 V3.0 基础上"追加"数据：
  * 帧结构：始终使用 V3.0 基底（标准协议帧）
  * 命令字：保留 V3.0 全部 21 条标准命令，仅追加 Word 中新增的命令
  * 枚举表：保留 V3.0 标准枚举，追加 Word 中的新枚举/新取值
  * 属性表：以 V3.0 基底为底，Word 中定义的属性覆盖/追加

使用：
    from protocol_parser.docx_importer import import_from_docx
    cfg = import_from_docx("协议.docx")
    # 保存为 JSON
    import json
    with open("product/my_product.json", "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)
"""
from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ImporterError(Exception):
    """文档导入错误。"""


# ---------- 工具函数 ----------

def _normalize(s: str) -> str:
    """标准化字符串：去除空白、换行、特殊字符。"""
    if not s:
        return ""
    return re.sub(r"\s+", " ", s).strip()


def _find_int(s: str) -> int | None:
    """从字符串中提取整数（支持 0x 前缀）。"""
    if not s:
        return None
    s = s.strip()
    # 优先匹配 0x 开头
    m = re.search(r"0[xX]([0-9a-fA-F]+)", s)
    if m:
        return int(m.group(1), 16)
    # 匹配纯数字
    m = re.search(r"\b(\d+)\b", s)
    if m:
        return int(m.group(1))
    return None


def _find_hex_int(s: str) -> int | None:
    """从字符串中提取 hex 整数（支持 0A / 0E / 00 这种无前缀形式）。"""
    if not s:
        return None
    s = s.strip()
    # 0x 前缀
    m = re.search(r"0[xX]([0-9a-fA-F]+)", s)
    if m:
        return int(m.group(1), 16)
    # 纯 hex 字符串（1-8 位，含字母才算 hex，避免与十进制冲突）
    m = re.match(r"^[0-9a-fA-F]{1,8}$", s)
    if m:
        return int(s, 16)
    # 字符串开头的 hex 字段
    m = re.match(r"^([0-9a-fA-F]{1,8})\b", s)
    if m:
        candidate = m.group(1)
        # 必须含字母才认定是 hex（纯数字走十进制）
        if any(c in "abcdefABCDEF" for c in candidate):
            return int(candidate, 16)
        return int(candidate)
    return None


def _find_hex_bytes(s: str) -> str | None:
    """从字符串中提取 hex 字节序列，返回标准格式 '0xXXXX'。"""
    if not s:
        return None
    # 匹配 0xA5A5 / 0xA5 0xA5 / A5 A5 A5
    m = re.search(r"0[xX]([0-9a-fA-F]{2,4})", s)
    if m:
        return f"0x{m.group(1).upper()}"
    # 匹配 A5A5 这种连续 hex
    m = re.search(r"\b([0-9a-fA-F]{4})\b", s)
    if m:
        return f"0x{m.group(1).upper()}"
    return None


def _table_to_rows(table) -> list[list[str]]:
    """把 docx 表格转成二维字符串列表。"""
    rows = []
    for row in table.rows:
        cells = [_normalize(c.text) for c in row.cells]
        rows.append(cells)
    return rows


def _find_header_row(rows: list[list[str]], keywords: list[str]) -> int:
    """在表格中找表头行（包含所有关键词的行）。"""
    for i, row in enumerate(rows):
        row_text = " ".join(row).lower()
        if all(k.lower() in row_text for k in keywords):
            return i
    return -1


def _column_map(header: list[str], field_names: list[list[str]]) -> dict[str, int]:
    """根据表头建立字段名 → 列索引的映射。

    field_names 是 [[可能名1, 可能名2], ...] 形式。
    """
    mapping: dict[str, int] = {}
    header_lower = [h.lower() for h in header]
    for field, names in field_names.items():
        for i, h in enumerate(header_lower):
            if any(n.lower() in h for n in names):
                mapping[field] = i
                break
    return mapping


# ---------- 文档结构识别 ----------

@dataclass
class ParsedDocument:
    """从 Word 解析出的中间结构。"""
    product_name: str = ""
    description: str = ""
    frame_config: dict = field(default_factory=dict)
    commands: list[dict] = field(default_factory=list)
    attributes: dict = field(default_factory=dict)
    enums: dict = field(default_factory=dict)
    raw_tables: list[list[list[str]]] = field(default_factory=list)
    raw_paragraphs: list[str] = field(default_factory=list)


def _read_docx(path: str | Path) -> ParsedDocument:
    """读取 Word 文档，提取所有段落和表格。"""
    if not HAS_DOCX:
        raise ImporterError(
            "python-docx 未安装。请在命令行执行：\n"
            "  pip install python-docx -i https://pypi.tuna.tsinghua.edu.cn/simple"
        )
    p = Path(path)
    if not p.exists():
        raise ImporterError(f"文件不存在: {p}")

    doc = Document(p)
    parsed = ParsedDocument()

    # 段落
    for para in doc.paragraphs:
        text = _normalize(para.text)
        if text:
            parsed.raw_paragraphs.append(text)

    # 表格
    for table in doc.tables:
        rows = _table_to_rows(table)
        parsed.raw_tables.append(rows)

    # 从段落提取产品名/描述
    for text in parsed.raw_paragraphs[:20]:
        if any(k in text.lower() for k in ["协议", "protocol"]):
            if not parsed.product_name:
                parsed.product_name = text[:60]
                continue
        if not parsed.description and len(text) > 10:
            parsed.description = text[:200]
            break

    return parsed


# ---------- 帧结构识别 ----------

def _parse_frame_config(parsed: ParsedDocument) -> dict:
    """从文档识别帧结构配置。"""
    frame = {
        "header": "0xA5A5",
        "header_size": 2,
        "ver": "0x03",
        "ver_offset": 2,
        "ver_size": 1,
        "cmd_offset": 3,
        "length_offset": 4,
        "length_size": 2,
        "length_byte_order": "big",
        "checksum": {
            "algorithm": "sum",
            "length": 1,
            "covers": "from_start_to_checksum_exclusive",
        },
    }

    # 从段落中找帧头、版本、校验等关键词
    all_text = "\n".join(parsed.raw_paragraphs)
    for table in parsed.raw_tables:
        for row in table:
            all_text += "\n" + " ".join(row)

    # 帧头：寻找 "帧头 0xA5A5" / "起始符 0xA5" 等
    m = re.search(r"(?:帧头|起始符|header|frame header)[^\n]*?(0[xX][0-9a-fA-F]{2,4})", all_text, re.IGNORECASE)
    if m:
        val = m.group(1)
        frame["header"] = val
        hex_str = val[2:]
        frame["header_size"] = len(hex_str) // 2

    # 版本
    m = re.search(r"(?:版本|version|ver)[^\n]*?(0[xX][0-9a-fA-F]+)", all_text, re.IGNORECASE)
    if m:
        frame["ver"] = m.group(1)

    # 校验算法
    if re.search(r"异或|XOR", all_text, re.IGNORECASE):
        frame["checksum"]["algorithm"] = "xor"
    elif re.search(r" CRC", all_text, re.IGNORECASE):
        frame["checksum"]["algorithm"] = "crc8"
    elif re.search(r"求和|sum|累加", all_text, re.IGNORECASE):
        frame["checksum"]["algorithm"] = "sum"

    return frame


# ---------- 命令表识别 ----------

CMD_FIELD_NAMES = {
    "cmd_code": ["cmd", "命令字", "命令码", "命令", "数值", "cmd_code", "value"],
    "name": ["name", "名称", "命令名", "功能", "描述", "description"],
    "description": ["说明", "描述", "功能描述", "备注", "description"],
    "direction": ["方向", "direction", "类型"],
    "format": ["格式", "format", "数据格式"],
}


def _parse_commands(parsed: ParsedDocument) -> list[dict]:
    """识别命令字表。

    策略：找到含 "命令字"/"cmd"/"数值" 等表头的表格，按列解析。
    """
    commands: list[dict] = []
    seen_codes: set[int] = set()

    for table in parsed.raw_tables:
        if not table:
            continue
        # 找表头行
        header_idx = _find_header_row(table, ["命令字"])
        if header_idx < 0:
            header_idx = _find_header_row(table, ["cmd"])
        if header_idx < 0:
            # V3.0 风格：Name + 方向 + 数值 + 说明
            for i, row in enumerate(table):
                row_text = " ".join(row).lower()
                if ("数值" in row_text or "value" in row_text) and ("name" in row_text or "名称" in row_text or "说明" in row_text):
                    header_idx = i
                    break
        if header_idx < 0:
            # 也尝试 "命令" + "名称" 组合
            for i, row in enumerate(table):
                row_text = " ".join(row).lower()
                if ("命令" in row_text or "cmd" in row_text) and ("名称" in row_text or "name" in row_text or "功能" in row_text):
                    header_idx = i
                    break
        if header_idx < 0:
            continue

        header = table[header_idx]
        col_map = _column_map(header, CMD_FIELD_NAMES)
        if "cmd_code" not in col_map:
            continue

        cmd_col = col_map["cmd_code"]
        name_col = col_map.get("name")
        desc_col = col_map.get("description", name_col)
        dir_col = col_map.get("direction")
        fmt_col = col_map.get("format")

        for row in table[header_idx + 1:]:
            if cmd_col >= len(row):
                continue
            cmd_text = row[cmd_col]
            # 命令字通常以 hex 形式书写（0x20 / 20 / 0xA5 等），优先按 hex 解析
            cmd_code = _find_hex_int(cmd_text)
            if cmd_code is None:
                cmd_code = _find_int(cmd_text)
            if cmd_code is None or cmd_code in seen_codes:
                continue

            name = row[name_col] if name_col is not None and name_col < len(row) else ""
            desc = row[desc_col] if desc_col is not None and desc_col < len(row) else ""
            direction = row[dir_col] if dir_col is not None and dir_col < len(row) else ""
            fmt = row[fmt_col] if fmt_col is not None and fmt_col < len(row) else ""

            # 推断 format
            if not fmt:
                fmt = _infer_format_from_text(name + " " + desc)

            cmd = {
                "cmd_code": f"0x{cmd_code:02X}",
                "name": name or f"cmd_{cmd_code:02X}",
                "description": desc,
            }
            if direction:
                cmd["direction"] = direction

            # 双向命令（同 cmd_code 有 request/response）
            # 简化处理：根据方向字段分拆；无方向字段则默认 request/response 都用 attr_list
            if direction.lower() in ("请求", "request", "下行", "命令下发"):
                cmd["request"] = {"format": fmt or "raw", "name": name}
                cmd["response"] = {"format": "raw", "name": name + " 响应"}
            elif direction.lower() in ("响应", "response", "上行", "应答"):
                # 已存在同 cmd_code 的 request 则补充；否则当作双向
                existing = next((c for c in commands if c.get("cmd_code") == f"0x{cmd_code:02X}"), None)
                if existing and "request" in existing:
                    existing["response"] = {"format": fmt or "raw", "name": name}
                    continue
                else:
                    cmd["request"] = {"format": "raw", "name": name + " 请求"}
                    cmd["response"] = {"format": fmt or "raw", "name": name}
            else:
                # 无方向：默认双向
                cmd["request"] = {"format": "raw", "name": "请求"}
                cmd["response"] = {"format": fmt or "attr_list", "name": "响应"}

            commands.append(cmd)
            seen_codes.add(cmd_code)

    return commands


def _infer_format_from_text(text: str) -> str:
    """根据命令描述推断数据格式。"""
    text_lower = text.lower()
    if "心跳" in text or "heartbeat" in text_lower:
        return "module_status"
    if "设备信息" in text or "version" in text_lower or "dev_info" in text_lower:
        return "dev_version"
    if "配网" in text or "net_config" in text_lower:
        return "net_config"
    if "时间" in text and "get" in text_lower:
        return "get_time_resp"
    if "快照" in text or "snapshot" in text_lower:
        return "attr_list"
    if "上报" in text or "report" in text_lower:
        return "attr_list"
    if "事件" in text or "event" in text_lower:
        return "event"
    if "行为" in text or "action" in text_lower:
        return "msg_id_then_action"
    if "ota" in text_lower or "升级" in text:
        return "raw"
    if "错误" in text or "errcode" in text_lower:
        return "errcode"
    return "attr_list"


# ---------- 属性表识别 ----------

ATTR_FIELD_NAMES = {
    "attrid": ["attrid", "attr id", "属性id", "属性id", "属性", "aid"],
    "name": ["name", "名称", "属性名", "属性名称"],
    "typeid": ["typeid", "type id", "类型id", "类型", "type"],
    "access": ["access", "权限", "访问", "读写"],
    "unit": ["unit", "单位"],
    "range": ["range", "范围", "取值"],
    "enum": ["enum", "枚举", "取值说明", "说明"],
}


TYPEID_KEYWORDS = {
    "bool": 0, "布尔": 0,
    "int8": 1,
    "uint8": 2, "u8": 2,
    "int16": 3, "i16": 3,
    "uint16": 4, "u16": 4,
    "int32": 5, "i32": 5,
    "uint32": 6, "u32": 6,
    "int64": 7,
    "uint64": 8,
    "float": 9, "float32": 9,
    "string": 11, "字符串": 11,
    "f1_u16": 15, "f1u16": 15,
    "f2_u16": 16, "f2u16": 16,
}


def _parse_typeid(text: str) -> int | None:
    """从文本解析 typeid。"""
    if not text:
        return None
    text_lower = text.strip().lower()
    # 直接是数字
    n = _find_int(text)
    if n is not None and 0 <= n <= 24:
        return n
    # 关键词匹配
    for kw, tid in TYPEID_KEYWORDS.items():
        if kw in text_lower:
            return tid
    return None


def _parse_attributes(parsed: ParsedDocument) -> dict:
    """识别属性表。"""
    attributes: dict = {}

    for table in parsed.raw_tables:
        if not table:
            continue
        # 找表头
        header_idx = -1
        for i, row in enumerate(table):
            row_text = " ".join(row).lower()
            if ("attrid" in row_text or "属性id" in row_text or "属性 id" in row_text) and ("name" in row_text or "名称" in row_text):
                header_idx = i
                break
        if header_idx < 0:
            continue

        header = table[header_idx]
        col_map = _column_map(header, ATTR_FIELD_NAMES)
        if "attrid" not in col_map:
            continue

        aid_col = col_map["attrid"]
        name_col = col_map.get("name")
        type_col = col_map.get("typeid")
        access_col = col_map.get("access")
        unit_col = col_map.get("unit")
        range_col = col_map.get("range")
        enum_col = col_map.get("enum")

        for row in table[header_idx + 1:]:
            if aid_col >= len(row):
                continue
            aid_text = row[aid_col]
            # 属性 ID 通常以 hex 形式书写（0A / 0E / E0 等），优先按 hex 解析
            aid = _find_hex_int(aid_text)
            if aid is None:
                aid = _find_int(aid_text)
            if aid is None:
                continue

            key = f"0x{aid:02X}"
            if key in attributes:
                continue

            attr: dict[str, Any] = {
                "name": row[name_col] if name_col is not None and name_col < len(row) else "",
            }

            # typeid
            if type_col is not None and type_col < len(row):
                tid = _parse_typeid(row[type_col])
                if tid is not None:
                    attr["typeid"] = tid

            # access
            if access_col is not None and access_col < len(row):
                attr["access"] = row[access_col]

            # unit
            if unit_col is not None and unit_col < len(row) and row[unit_col]:
                attr["unit"] = row[unit_col]

            # range
            if range_col is not None and range_col < len(row) and row[range_col]:
                attr["range"] = row[range_col]

            # enum（取值说明列）
            if enum_col is not None and enum_col < len(row) and row[enum_col]:
                enum_text = row[enum_col]
                enum_map = _parse_enum_text(enum_text)
                if enum_map:
                    attr["enum"] = enum_map

            attributes[key] = attr

    return attributes


def _parse_enum_text(text: str) -> dict[str, str]:
    """解析枚举文本，如 "0:关 1:开" / "0=关闭, 1=打开"。"""
    if not text:
        return {}
    result: dict[str, str] = {}
    # 分割多个枚举项
    parts = re.split(r"[;,，；\n]+", text)
    for part in parts:
        # 匹配 "数字:文本" / "数字=文本" / "数字 文本"
        m = re.match(r"\s*(\d+)\s*[:=：]\s*(.+)", part)
        if m:
            k, v = m.group(1), m.group(2).strip()
            if v:
                result[k] = v
    return result


# ---------- V3.0 基底协议 ----------

def _resource_path(relative: str) -> Path:
    """获取资源路径，兼容开发模式和 PyInstaller 打包模式。"""
    if hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS) / relative
    base = Path(__file__).resolve().parent
    candidate = base / relative
    if candidate.exists():
        return candidate
    return base.parent / relative


def _find_v3_base_file() -> Path | None:
    """查找 V3.0 基底协议文件 v3_serial.json。

    查找顺序：
    1. PyInstaller 打包资源：sys._MEIPASS/product/v3_serial.json
    2. 开发模式：项目根目录/product/v3_serial.json
    3. exe 同目录：exe_dir/product/v3_serial.json
    """
    candidates = [
        _resource_path("product") / "v3_serial.json",
    ]
    # exe 同目录（用户可见的 product 目录）
    if getattr(sys, "frozen", False):
        candidates.append(Path(sys.executable).resolve().parent / "product" / "v3_serial.json")
    for c in candidates:
        if c.exists():
            return c
    return None


def load_v3_base() -> dict | None:
    """加载 V3.0 基底协议。找不到时返回 None（退化为旧行为）。"""
    p = _find_v3_base_file()
    if p is None:
        return None
    try:
        with p.open("r", encoding="utf-8") as f:
            cfg = json.load(f)
        if "commands" in cfg and "frame" in cfg:
            return cfg
    except (OSError, json.JSONDecodeError):
        pass
    return None


def _parse_cmd_code_int(cmd_code: Any) -> int | None:
    """把 cmd_code（如 '0x20' / 32）解析为 int。"""
    if isinstance(cmd_code, int):
        return cmd_code
    if isinstance(cmd_code, str):
        s = cmd_code.strip().lower()
        try:
            return int(s, 16) if s.startswith("0x") else int(s, 0)
        except ValueError:
            return None
    return None


def merge_with_v3_base(imported: dict, base: dict) -> dict:
    """把 Word 导入的配置合并到 V3.0 基底之上。

    合并规则：
    - frame: 始终使用 V3.0 基底（标准协议帧结构不可被覆盖）
    - commands: V3.0 基底全部保留，追加 Word 中新增命令（同 cmd_code 不覆盖基底）
    - enums: V3.0 基底全部保留；同枚举名合并取值（基底已有则保留），新枚举直接追加
    - attributes: V3.0 基底为底，Word 中定义的属性覆盖同名 attrid（产品属性以 Word 为准）
    """
    base_commands = base.get("commands", []) or []
    base_enums = dict(base.get("enums", {}) or {})
    base_attrs = dict(base.get("attributes", {}) or {})
    base_frame = base.get("frame", {}) or {}

    # 命令：基底优先，追加 Word 新增
    seen_codes: set[int] = set()
    merged_commands: list[dict] = []
    for cmd in base_commands:
        code = _parse_cmd_code_int(cmd.get("cmd_code"))
        if code is not None:
            seen_codes.add(code)
        merged_commands.append(cmd)
    for cmd in imported.get("commands", []) or []:
        code = _parse_cmd_code_int(cmd.get("cmd_code"))
        if code is None or code in seen_codes:
            continue
        seen_codes.add(code)
        merged_commands.append(cmd)

    # 枚举：基底优先，合并新取值
    merged_enums: dict[str, dict] = {k: dict(v) for k, v in base_enums.items()}
    for enum_name, enum_map in (imported.get("enums", {}) or {}).items():
        if not isinstance(enum_map, dict):
            continue
        if enum_name in merged_enums:
            for k, v in enum_map.items():
                if k not in merged_enums[enum_name]:
                    merged_enums[enum_name][k] = v
        else:
            merged_enums[enum_name] = dict(enum_map)

    # 属性：基底为底，Word 覆盖同名 attrid
    merged_attrs: dict[str, dict] = {k: dict(v) for k, v in base_attrs.items()}
    for aid, attr in (imported.get("attributes", {}) or {}).items():
        if not isinstance(attr, dict):
            continue
        if aid in merged_attrs:
            merged = dict(merged_attrs[aid])
            merged.update(attr)
            merged_attrs[aid] = merged
        else:
            merged_attrs[aid] = dict(attr)

    return {
        "product": imported.get("product", "product"),
        "description": imported.get("description", ""),
        "version": "3.0",
        "base_protocol": "v3_serial",
        "frame": base_frame,
        "enums": merged_enums,
        "commands": merged_commands,
        "attributes": merged_attrs,
        "_imported_from": imported.get("_imported_from", ""),
    }


# ---------- 主入口 ----------

def import_from_docx(path: str | Path, product_name: str | None = None) -> dict:
    """从 Word 文档生成协议 JSON 配置。

    策略：以 V3.0 串口标准协议为基底，Word 文档中的命令/属性/枚举仅在基底之上追加。
    这样可保证导入新产品后，V3.0 的 21 条标准命令（心跳、查询、配网等）依然可用，
    不会出现"未知命令"。

    Args:
        path: Word 文档路径
        product_name: 自定义产品名（不指定则从文档提取）

    Returns:
        协议配置字典，可直接保存为 JSON
    """
    parsed = _read_docx(path)

    # 进一步解析 Word 文档内容
    parsed.frame_config = _parse_frame_config(parsed)
    parsed.commands = _parse_commands(parsed)
    parsed.attributes = _parse_attributes(parsed)

    # 产品名
    name = product_name or parsed.product_name or Path(path).stem
    # 清理产品名（只保留字母数字下划线）
    name = re.sub(r"[^\w\u4e00-\u9fff]", "_", name)
    name = re.sub(r"_+", "_", name).strip("_").lower() or "product"

    word_cfg = {
        "product": name,
        "description": parsed.description or f"从 {Path(path).name} 导入的协议（基于 V3.0 基底）",
        "version": "3.0",
        "frame": parsed.frame_config,
        "enums": parsed.enums,
        "commands": parsed.commands,
        "attributes": parsed.attributes,
        "_imported_from": str(Path(path).name),
    }

    # 合并到 V3.0 基底之上（基底找不到时退化为仅使用 Word 内容）
    base = load_v3_base()
    if base is not None:
        return merge_with_v3_base(word_cfg, base)
    return word_cfg


def save_protocol_json(cfg: dict, output_path: str | Path) -> Path:
    """保存协议配置为 JSON 文件。"""
    p = Path(output_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    with p.open("w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)
    return p


def import_and_save(docx_path: str | Path, protocols_dir: str | Path, product_name: str | None = None) -> tuple[dict, Path]:
    """导入 Word 文档并保存为 JSON。

    Returns: (配置字典, 保存路径)
    """
    cfg = import_from_docx(docx_path, product_name)
    product = cfg["product"]
    out_path = Path(protocols_dir) / f"{product}.json"
    save_protocol_json(cfg, out_path)
    return cfg, out_path


def check_docx_available() -> bool:
    """检查 python-docx 是否可用。"""
    return HAS_DOCX
